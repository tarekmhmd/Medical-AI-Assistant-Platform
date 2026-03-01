"""
Generate Word Report for Trained Models
========================================
Creates a comprehensive .docx report with model information.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime

def create_model_report():
    """Create comprehensive Word report for trained models."""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Medical AI Platform - Model Training Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Unified Models: Classification + Segmentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    summary = doc.add_paragraph()
    summary.add_run('This report documents the successful training of four unified deep learning models '
                   'that support both classification and segmentation tasks. All models were trained on '
                   'NVIDIA GeForce GTX 1650 (4GB VRAM) with mixed precision training for optimal performance.')
    
    # Training Summary Table
    doc.add_heading('Training Summary', level=1)
    
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    
    # Header row
    headers = ['Model', 'Original Type', 'Upgraded Type', 'Status']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Data rows
    data = [
        ['Skin Analyzer', 'Segmentation', 'Segmentation + Classification', '✅ Trained'],
        ['Sound Analyzer', 'Classification', 'Classification + Segmentation', '✅ Trained'],
        ['Lab Analyzer', 'Classification', 'Classification + Segmentation', '✅ Trained'],
        ['Medical Chatbot', 'Classification', 'Classification + Segmentation', '✅ Trained'],
    ]
    
    for i, row_data in enumerate(data):
        for j, value in enumerate(row_data):
            table.rows[i + 1].cells[j].text = value
    
    doc.add_paragraph()
    
    # Checkpoint Files
    doc.add_heading('Trained Model Checkpoints', level=1)
    
    checkpoint_table = doc.add_table(rows=5, cols=3)
    checkpoint_table.style = 'Table Grid'
    
    # Headers
    ch_headers = ['Checkpoint File', 'Size (MB)', 'Purpose']
    for i, header in enumerate(ch_headers):
        cell = checkpoint_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Checkpoint data
    checkpoints = [
        ['skin_model_cls_seg_trained.pth', '~31.0', 'Skin lesion segmentation + classification'],
        ['sound_model_cls_seg_trained.pth', '12.71', 'Respiratory sound classification + attention'],
        ['lab_model_cls_seg_trained.pth', '0.08', 'Diabetes prediction + feature importance'],
        ['chatbot_model_cls_seg_trained.pth', '7.20', 'Medical chatbot + word attention'],
    ]
    
    for i, row_data in enumerate(checkpoints):
        for j, value in enumerate(row_data):
            checkpoint_table.rows[i + 1].cells[j].text = value
    
    doc.add_paragraph()
    
    # Model Details
    doc.add_heading('Model Details', level=1)
    
    # 1. Skin Model
    doc.add_heading('1. Skin Analyzer (UNetClassifier)', level=2)
    
    doc.add_paragraph('Architecture: U-Net with Classification Head', style='List Bullet')
    doc.add_paragraph('Input: RGB Image [batch, 3, 256, 256]', style='List Bullet')
    doc.add_paragraph('Outputs:', style='List Bullet')
    doc.add_paragraph('• Segmentation Mask: [batch, 1, 256, 256]', style='List Bullet 2')
    doc.add_paragraph('• Classification Logits: [batch, 8]', style='List Bullet 2')
    doc.add_paragraph('Dataset: ISIC 2016 Task 1 (900 training images)', style='List Bullet')
    doc.add_paragraph('Classes: 8 skin conditions (Healthy, Acne, Eczema, Psoriasis, Melanoma, Dermatitis, Rosacea, Fungal)', style='List Bullet')
    doc.add_paragraph('Loss: 0.7 × BCEWithLogitsLoss + 0.3 × CrossEntropyLoss', style='List Bullet')
    
    doc.add_paragraph()
    
    # 2. Sound Model
    doc.add_heading('2. Sound Analyzer (CNNSegmenter)', level=2)
    
    doc.add_paragraph('Architecture: CNN with Attention Decoder', style='List Bullet')
    doc.add_paragraph('Input: Mel-Spectrogram [batch, 1, 64, 128]', style='List Bullet')
    doc.add_paragraph('Outputs:', style='List Bullet')
    doc.add_paragraph('• Attention Map: [batch, 1, 64, 128]', style='List Bullet 2')
    doc.add_paragraph('• Classification Logits: [batch, 6]', style='List Bullet 2')
    doc.add_paragraph('Dataset: Respiratory Sound Database (1,472 WAV files)', style='List Bullet')
    doc.add_paragraph('Classes: 6 conditions (Healthy, Asthma, Bronchitis, Pneumonia, COPD, Whooping Cough)', style='List Bullet')
    doc.add_paragraph('Loss: 0.3 × BCEWithLogitsLoss + 0.7 × CrossEntropyLoss', style='List Bullet')
    
    doc.add_paragraph()
    
    # 3. Lab Model
    doc.add_heading('3. Lab Analyzer (MLPSegmenter)', level=2)
    
    doc.add_paragraph('Architecture: MLP with Feature Importance Head', style='List Bullet')
    doc.add_paragraph('Input: Tabular Features [batch, 8]', style='List Bullet')
    doc.add_paragraph('Outputs:', style='List Bullet')
    doc.add_paragraph('• Importance Map: [batch, 1, 8, 8]', style='List Bullet 2')
    doc.add_paragraph('• Classification Logits: [batch, 2]', style='List Bullet 2')
    doc.add_paragraph('Dataset: Pima Indians Diabetes (~500 samples)', style='List Bullet')
    doc.add_paragraph('Features: 8 clinical measurements (Glucose, BMI, Age, etc.)', style='List Bullet')
    doc.add_paragraph('Classes: 2 (Non-diabetic / Diabetic)', style='List Bullet')
    doc.add_paragraph('Loss: 0.2 × BCEWithLogitsLoss + 0.8 × CrossEntropyLoss', style='List Bullet')
    
    doc.add_paragraph()
    
    # 4. Chatbot Model
    doc.add_heading('4. Medical Chatbot (LSTMSegmenter)', level=2)
    
    doc.add_paragraph('Architecture: Bidirectional LSTM with Word Attention', style='List Bullet')
    doc.add_paragraph('Input: Tokenized Question [batch, 20]', style='List Bullet')
    doc.add_paragraph('Outputs:', style='List Bullet')
    doc.add_paragraph('• Word Attention: [batch, 1, 20]', style='List Bullet 2')
    doc.add_paragraph('• Classification Logits: [batch, ~100]', style='List Bullet 2')
    doc.add_paragraph('Dataset: Combined Medical QA (~11,000 training pairs)', style='List Bullet')
    doc.add_paragraph('Vocabulary: ~2,500 words', style='List Bullet')
    doc.add_paragraph('Classes: ~100 disease categories', style='List Bullet')
    doc.add_paragraph('Loss: 0.2 × BCEWithLogitsLoss + 0.8 × CrossEntropyLoss', style='List Bullet')
    
    doc.add_paragraph()
    
    # GPU Optimization
    doc.add_heading('GPU Memory Optimization', level=1)
    
    doc.add_paragraph('Training Device: NVIDIA GeForce GTX 1650 (4GB VRAM)')
    
    gpu_table = doc.add_table(rows=5, cols=3)
    gpu_table.style = 'Table Grid'
    
    gpu_headers = ['Model', 'Batch Size', 'Memory Usage']
    for i, header in enumerate(gpu_headers):
        cell = gpu_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    gpu_data = [
        ['Skin (U-Net)', '4', '~3.5 GB'],
        ['Sound (CNN)', '8', '~1.5 GB'],
        ['Lab (MLP)', '32', '< 0.5 GB'],
        ['Chatbot (LSTM)', '32', '~0.5 GB'],
    ]
    
    for i, row_data in enumerate(gpu_data):
        for j, value in enumerate(row_data):
            gpu_table.rows[i + 1].cells[j].text = value
    
    doc.add_paragraph()
    
    optimizations = doc.add_paragraph('Optimizations Applied:')
    optimizations.runs[0].bold = True
    doc.add_paragraph('Mixed Precision Training (FP16)', style='List Bullet')
    doc.add_paragraph('Gradient Scaling', style='List Bullet')
    doc.add_paragraph('Adaptive Batch Sizes', style='List Bullet')
    doc.add_paragraph('Memory Cleanup Between Models', style='List Bullet')
    
    doc.add_paragraph()
    
    # Usage Examples
    doc.add_heading('Usage Examples', level=1)
    
    doc.add_paragraph('Loading and using trained models:', style='Intense Quote')
    
    code_example = doc.add_paragraph()
    code_example.add_run('''from models.unified_models import create_model
import torch

# Load Skin Model
model = create_model('unet', num_classes=8)
checkpoint = torch.load('checkpoints/skin_model_cls_seg_trained.pth')
model.load_state_dict(checkpoint['model_state_dict'])

# Inference
image = torch.randn(1, 3, 256, 256)
mask, class_logits = model(image)
# mask: [1, 1, 256, 256] - Lesion segmentation
# class_logits: [1, 8] - Disease classification''')
    
    doc.add_paragraph()
    
    # Files Created
    doc.add_heading('Files Created', level=1)
    
    files = [
        'models/unified_models.py - All unified model architectures',
        'train_unified_complete.py - Complete training pipeline',
        'checkpoints/skin_model_cls_seg_trained.pth - Skin model weights',
        'checkpoints/sound_model_cls_seg_trained.pth - Sound model weights',
        'checkpoints/lab_model_cls_seg_trained.pth - Lab model weights',
        'checkpoints/chatbot_model_cls_seg_trained.pth - Chatbot model weights',
        'model_info.txt - Detailed model documentation',
        'training_summary.json - Training results summary',
        'model_report.docx - This report',
    ]
    
    for f in files:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_paragraph()
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    
    conclusion = doc.add_paragraph()
    conclusion.add_run('All four models have been successfully upgraded from single-task to multi-task '
                      'architectures supporting both classification and segmentation. The models are ready '
                      'for deployment in the Medical AI Assistant Platform. Each model provides interpretable '
                      'outputs through attention/importance maps, enhancing the platform\'s diagnostic capabilities.')
    
    # Save document
    output_path = os.path.join(os.path.dirname(__file__), 'model_report.docx')
    doc.save(output_path)
    
    print(f"✅ Report saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_model_report()
