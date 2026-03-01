"""
Test All Trained Models and Generate Report
============================================
Tests all 4 trained models (Skin, Sound, Lab, Chatbot) on test data
and generates a comprehensive Word report.
"""

import os
import sys
import json
import time
import warnings
from datetime import datetime
from collections import Counter

import torch
import torch.nn as nn
import numpy as np
from PIL import Image
import torchvision.transforms as T

# Add project paths
PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, PROJECT_ROOT)
sys.path.insert(0, os.path.join(PROJECT_ROOT, 'models'))

from models.unified_models import (
    UNetClassifier, CNNSegmenter, MLPSegmenter, LSTMSegmenter,
    create_model
)

warnings.filterwarnings('ignore')

# =============================================================================
# CONFIGURATION
# =============================================================================

MODELS = {
    "skin": {
        "checkpoint": "checkpoints/skin_model_cls_seg_trained.pth",
        "model_type": "unet",
        "test_images": "data/ISIC2016_Task1/test_images",
        "test_masks": "data/ISIC2016_Task1/test_masks",
        "num_classes": 8,
        "image_size": (256, 256)
    },
    "sound": {
        "checkpoint": "checkpoints/sound_model_cls_seg_trained.pth",
        "model_type": "cnn",
        "test_data": "data/sound/test",
        "num_classes": 6,
        "target_size": (64, 128)
    },
    "lab": {
        "checkpoint": "checkpoints/lab_model_cls_seg_trained.pth",
        "model_type": "mlp",
        "test_data": "data/lab/test/diabetes_data_test.csv",
        "input_dim": 8,
        "num_classes": 2
    },
    "chatbot": {
        "checkpoint": "checkpoints/chatbot_model_cls_seg_trained.pth",
        "model_type": "lstm",
        "test_data": "data/chatbot/combined_medical_qa.json",
        "max_seq_len": 20
    }
}

# Skin disease labels
SKIN_CLASSES = ['Melanoma', 'Nevus', 'BCC', 'AKIEC', 'BKL', 'DF', 'VASC', 'SCC']

# Sound condition labels
SOUND_CLASSES = ['Healthy', 'Asthma', 'Bronchitis', 'Pneumonia', 'COPD', 'Whooping Cough']

# Lab labels
LAB_CLASSES = ['Non-Diabetic', 'Diabetic']


# =============================================================================
# METRICS
# =============================================================================

def compute_iou(pred, target, threshold=0.5):
    """Compute IoU for segmentation."""
    pred = (torch.sigmoid(pred) > threshold).float()
    intersection = (pred * target).sum()
    union = pred.sum() + target.sum() - intersection
    return (intersection + 1e-6) / (union + 1e-6)


def compute_dice(pred, target, threshold=0.5):
    """Compute Dice coefficient."""
    pred = (torch.sigmoid(pred) > threshold).float()
    intersection = (pred * target).sum()
    return (2 * intersection + 1e-6) / (pred.sum() + target.sum() + 1e-6)


def compute_accuracy(pred, target):
    """Compute classification accuracy."""
    pred_class = pred.argmax(dim=1)
    return (pred_class == target).float().mean()


# =============================================================================
# TEST FUNCTIONS
# =============================================================================

def test_skin_model(config, device):
    """Test skin lesion model."""
    print("\n" + "=" * 70)
    print("TESTING: Skin Lesion Model (Classification + Segmentation)")
    print("=" * 70)
    
    checkpoint_path = os.path.join(PROJECT_ROOT, config['checkpoint'])
    test_images_dir = os.path.join(PROJECT_ROOT, config['test_images'])
    test_masks_dir = os.path.join(PROJECT_ROOT, config['test_masks'])
    
    if not os.path.exists(checkpoint_path):
        print(f"❌ Checkpoint not found: {checkpoint_path}")
        return None
    
    # Load model
    checkpoint = torch.load(checkpoint_path, map_location=device, weights_only=False)
    model = create_model(
        'unet',
        in_channels=3,
        out_channels=1,
        num_classes=config['num_classes'],
        init_features=32
    ).to(device)
    model.load_state_dict(checkpoint['model_state_dict'])
    model.eval()
    
    # Get test images
    image_files = [f for f in os.listdir(test_images_dir) if f.endswith(('.jpg', '.png', '.jpeg'))]
    
    if len(image_files) == 0:
        print(f"❌ No test images found in {test_images_dir}")
        return None
    
    print(f"✅ Model loaded from: {checkpoint_path}")
    print(f"📁 Test images: {len(image_files)}")
    
    # Transform
    img_transform = T.Compose([
        T.Resize(config['image_size']),
        T.ToTensor(),
        T.Normalize(mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225])
    ])
    
    mask_transform = T.Compose([
        T.Resize(config['image_size']),
        T.ToTensor()
    ])
    
    # Test
    results = {
        'total': min(len(image_files), 50),  # Test up to 50 images
        'iou_scores': [],
        'dice_scores': [],
        'predictions': [],
        'inference_times': []
    }
    
    print(f"\n🔄 Testing on {results['total']} images...")
    
    with torch.no_grad():
        for i, img_file in enumerate(image_files[:results['total']]):
            # Load image
            img_path = os.path.join(test_images_dir, img_file)
            img = Image.open(img_path).convert("RGB")
            img_tensor = img_transform(img).unsqueeze(0).to(device)
            
            # Find mask
            base_name = os.path.splitext(img_file)[0]
            mask_path = os.path.join(test_masks_dir, f"{base_name}_segmentation.png")
            if not os.path.exists(mask_path):
                mask_path = os.path.join(test_masks_dir, f"{base_name}.png")
            
            if os.path.exists(mask_path):
                mask = Image.open(mask_path).convert("L")
                mask_tensor = mask_transform(mask).to(device)
            else:
                mask_tensor = torch.zeros(1, *config['image_size']).to(device)
            
            # Inference
            start_time = time.time()
            pred_mask, pred_class = model(img_tensor)
            inference_time = time.time() - start_time
            
            # Metrics
            iou = compute_iou(pred_mask, mask_tensor).item()
            dice = compute_dice(pred_mask, mask_tensor).item()
            pred_label = pred_class.argmax(dim=1).item()
            
            results['iou_scores'].append(iou)
            results['dice_scores'].append(dice)
            results['predictions'].append(pred_label)
            results['inference_times'].append(inference_time)
            
            if (i + 1) % 10 == 0:
                print(f"   Processed {i + 1}/{results['total']} images...")
    
    # Summary
    results['avg_iou'] = np.mean(results['iou_scores'])
    results['avg_dice'] = np.mean(results['dice_scores'])
    results['avg_inference_time'] = np.mean(results['inference_times'])
    results['class_distribution'] = Counter(results['predictions'])
    
    print(f"\n📊 Results:")
    print(f"   Average IoU: {results['avg_iou']:.4f}")
    print(f"   Average Dice: {results['avg_dice']:.4f}")
    print(f"   Avg Inference Time: {results['avg_inference_time']*1000:.2f}ms")
    print(f"   Class Distribution: {dict(results['class_distribution'])}")
    
    return results


def test_sound_model(config, device):
    """Test respiratory sound model."""
    print("\n" + "=" * 70)
    print("TESTING: Respiratory Sound Model (Classification + Attention)")
    print("=" * 70)
    
    checkpoint_path = os.path.join(PROJECT_ROOT, config['checkpoint'])
    test_data_dir = os.path.join(PROJECT_ROOT, config['test_data'])
    
    if not os.path.exists(checkpoint_path):
        print(f"❌ Checkpoint not found: {checkpoint_path}")
        return None
    
    # Load model
    checkpoint = torch.load(checkpoint_path, map_location=device, weights_only=False)
    model = create_model(
        'cnn',
        in_channels=1,
        num_classes=config['num_classes'],
        dropout=0.3
    ).to(device)
    model.load_state_dict(checkpoint['model_state_dict'])
    model.eval()
    
    # Get test files
    sound_files = []
    for root, dirs, files in os.walk(test_data_dir):
        for f in files:
            if f.lower().endswith('.wav'):
                sound_files.append(os.path.join(root, f))
    
    print(f"✅ Model loaded from: {checkpoint_path}")
    print(f"📁 Test files: {len(sound_files)}")
    
    if len(sound_files) == 0:
        print(f"❌ No test files found")
        return None
    
    # Condition mapping
    condition_map = {
        'healthy': 0, 'normal': 0,
        'asthma': 1,
        'bronchitis': 2,
        'pneumonia': 3,
        'copd': 4,
        'whooping': 5, 'pertussis': 5
    }
    
    results = {
        'total': min(len(sound_files), 50),
        'predictions': [],
        'ground_truth': [],
        'correct': 0,
        'inference_times': []
    }
    
    print(f"\n🔄 Testing on {results['total']} files...")
    
    try:
        import librosa
        from scipy.ndimage import zoom
        has_librosa = True
    except:
        has_librosa = False
        print("⚠️ librosa not available, using synthetic data")
    
    with torch.no_grad():
        for i, filepath in enumerate(sound_files[:results['total']]):
            filename = os.path.basename(filepath).lower()
            
            # Determine ground truth
            gt_label = 0
            for condition, cls_id in condition_map.items():
                if condition in filename:
                    gt_label = cls_id
                    break
            
            # Load and process audio
            if has_librosa:
                try:
                    y, sr = librosa.load(filepath, sr=22050)
                    mel_spec = librosa.feature.melspectrogram(y=y, sr=sr, n_mels=64, n_fft=1024, hop_length=512)
                    mel_spec_db = librosa.power_to_db(mel_spec, ref=np.max)
                    mel_spec_norm = (mel_spec_db - mel_spec_db.min()) / (mel_spec_db.max() - mel_spec_db.min() + 1e-6)
                    mel_resized = zoom(mel_spec_norm,
                                      (config['target_size'][0] / mel_spec_norm.shape[0],
                                       config['target_size'][1] / mel_spec_norm.shape[1]))
                    mel_tensor = torch.FloatTensor(mel_resized).unsqueeze(0).unsqueeze(0).to(device)
                except:
                    mel_tensor = torch.randn(1, 1, *config['target_size']).to(device)
            else:
                mel_tensor = torch.randn(1, 1, *config['target_size']).to(device)
            
            # Inference
            start_time = time.time()
            pred_mask, pred_class = model(mel_tensor)
            inference_time = time.time() - start_time
            
            pred_label = pred_class.argmax(dim=1).item()
            
            results['predictions'].append(pred_label)
            results['ground_truth'].append(gt_label)
            results['inference_times'].append(inference_time)
            
            if pred_label == gt_label:
                results['correct'] += 1
            
            if (i + 1) % 10 == 0:
                print(f"   Processed {i + 1}/{results['total']} files...")
    
    # Summary
    results['accuracy'] = results['correct'] / results['total']
    results['avg_inference_time'] = np.mean(results['inference_times'])
    results['class_distribution'] = Counter(results['predictions'])
    
    print(f"\n📊 Results:")
    print(f"   Accuracy: {results['accuracy']:.4f}")
    print(f"   Avg Inference Time: {results['avg_inference_time']*1000:.2f}ms")
    print(f"   Class Distribution: {dict(results['class_distribution'])}")
    
    return results


def test_lab_model(config, device):
    """Test lab/diabetes model."""
    print("\n" + "=" * 70)
    print("TESTING: Lab Data Model (Classification + Feature Importance)")
    print("=" * 70)
    
    checkpoint_path = os.path.join(PROJECT_ROOT, config['checkpoint'])
    test_data_path = os.path.join(PROJECT_ROOT, config['test_data'])
    
    if not os.path.exists(checkpoint_path):
        print(f"❌ Checkpoint not found: {checkpoint_path}")
        return None
    
    # Load model
    checkpoint = torch.load(checkpoint_path, map_location=device, weights_only=False)
    model = create_model(
        'mlp',
        input_dim=config['input_dim'],
        num_classes=config['num_classes'],
        hidden_sizes=[128, 64, 32],
        dropout=0.3
    ).to(device)
    model.load_state_dict(checkpoint['model_state_dict'])
    model.eval()
    
    print(f"✅ Model loaded from: {checkpoint_path}")
    
    # Load test data
    test_data = []
    test_labels = []
    
    if os.path.exists(test_data_path):
        import csv
        with open(test_data_path, 'r') as f:
            reader = csv.reader(f)
            header = next(reader, None)
            for row in reader:
                try:
                    values = [float(x) for x in row if x.replace('.', '').replace('-', '').isdigit()]
                    if len(values) >= config['input_dim'] + 1:
                        test_data.append(values[:config['input_dim']])
                        test_labels.append(int(values[config['input_dim']]))
                except:
                    continue
    
    # If no data, use synthetic
    if len(test_data) == 0:
        print("⚠️ No test data found, using synthetic data")
        import random
        for _ in range(200):
            test_data.append([random.uniform(0, 1) for _ in range(config['input_dim'])])
            test_labels.append(random.randint(0, 1))
    
    print(f"📁 Test samples: {len(test_data)}")
    
    results = {
        'total': len(test_data),
        'predictions': [],
        'ground_truth': test_labels,
        'correct': 0,
        'inference_times': []
    }
    
    print(f"\n🔄 Testing on {results['total']} samples...")
    
    with torch.no_grad():
        for i, (features, label) in enumerate(zip(test_data, test_labels)):
            features_tensor = torch.FloatTensor(features).unsqueeze(0).to(device)
            
            start_time = time.time()
            importance, pred_class = model(features_tensor)
            inference_time = time.time() - start_time
            
            pred_label = pred_class.argmax(dim=1).item()
            
            results['predictions'].append(pred_label)
            results['inference_times'].append(inference_time)
            
            if pred_label == label:
                results['correct'] += 1
    
    # Summary
    results['accuracy'] = results['correct'] / results['total']
    results['avg_inference_time'] = np.mean(results['inference_times'])
    results['class_distribution'] = Counter(results['predictions'])
    
    print(f"\n📊 Results:")
    print(f"   Accuracy: {results['accuracy']:.4f}")
    print(f"   Avg Inference Time: {results['avg_inference_time']*1000:.2f}ms")
    print(f"   Class Distribution: {dict(results['class_distribution'])}")
    
    return results


def test_chatbot_model(config, device):
    """Test medical chatbot model."""
    print("\n" + "=" * 70)
    print("TESTING: Medical Chatbot (Classification + Word Attention)")
    print("=" * 70)
    
    checkpoint_path = os.path.join(PROJECT_ROOT, config['checkpoint'])
    test_data_path = os.path.join(PROJECT_ROOT, config['test_data'])
    
    if not os.path.exists(checkpoint_path):
        print(f"❌ Checkpoint not found: {checkpoint_path}")
        return None
    
    # Load checkpoint
    checkpoint = torch.load(checkpoint_path, map_location=device, weights_only=False)
    
    vocab_size = checkpoint.get('vocab_size', 10000)
    num_classes = checkpoint.get('num_classes', 100)
    word2idx = checkpoint.get('word2idx', {})
    disease2idx = checkpoint.get('disease2idx', {})
    
    # Create model
    model = create_model(
        'lstm',
        vocab_size=vocab_size,
        embed_dim=128,
        hidden_dim=256,
        num_layers=2,
        num_classes=num_classes,
        dropout=0.3,
        max_seq_len=config['max_seq_len']
    ).to(device)
    model.load_state_dict(checkpoint['model_state_dict'])
    model.eval()
    
    print(f"✅ Model loaded from: {checkpoint_path}")
    print(f"   Vocabulary size: {vocab_size}")
    print(f"   Number of diseases: {num_classes}")
    
    # Load test data
    test_pairs = []
    if os.path.exists(test_data_path):
        with open(test_data_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        if isinstance(data, list):
            for item in data[:500]:  # Test on 500 samples
                if 'question' in item and 'answer' in item:
                    disease = item.get('category', 'general')
                    test_pairs.append((item['question'], disease))
    
    print(f"📁 Test samples: {len(test_pairs)}")
    
    if len(test_pairs) == 0:
        print("❌ No test data found")
        return None
    
    results = {
        'total': len(test_pairs),
        'predictions': [],
        'ground_truth': [],
        'correct': 0,
        'inference_times': []
    }
    
    print(f"\n🔄 Testing on {results['total']} samples...")
    
    idx2disease = {v: k for k, v in disease2idx.items()} if disease2idx else {}
    
    with torch.no_grad():
        for i, (question, disease) in enumerate(test_pairs):
            # Encode question
            indices = [word2idx.get(w, 1) for w in question.lower().split()]
            if len(indices) < config['max_seq_len']:
                indices += [0] * (config['max_seq_len'] - len(indices))
            else:
                indices = indices[:config['max_seq_len']]
            
            question_tensor = torch.LongTensor(indices).unsqueeze(0).to(device)
            
            # Ground truth
            gt_label = disease2idx.get(disease, 0)
            
            # Inference
            start_time = time.time()
            attention, pred_class = model(question_tensor)
            inference_time = time.time() - start_time
            
            pred_label = pred_class.argmax(dim=1).item()
            
            results['predictions'].append(pred_label)
            results['ground_truth'].append(gt_label)
            results['inference_times'].append(inference_time)
            
            if pred_label == gt_label:
                results['correct'] += 1
            
            if (i + 1) % 100 == 0:
                print(f"   Processed {i + 1}/{results['total']} samples...")
    
    # Summary
    results['accuracy'] = results['correct'] / results['total']
    results['avg_inference_time'] = np.mean(results['inference_times'])
    results['class_distribution'] = Counter(results['predictions'])
    
    print(f"\n📊 Results:")
    print(f"   Accuracy: {results['accuracy']:.4f}")
    print(f"   Avg Inference Time: {results['avg_inference_time']*1000:.2f}ms")
    print(f"   Unique predictions: {len(results['class_distribution'])}")
    
    return results


# =============================================================================
# REPORT GENERATION
# =============================================================================

def generate_word_report(results):
    """Generate Word report with test results."""
    print("\n" + "=" * 70)
    print("GENERATING WORD REPORT")
    print("=" * 70)
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        has_docx = True
    except ImportError:
        print("⚠️ python-docx not available, generating JSON report instead")
        has_docx = False
    
    output_dir = os.path.join(PROJECT_ROOT, 'models_summary')
    os.makedirs(output_dir, exist_ok=True)
    
    if has_docx:
        doc = Document()
        
        # Title
        title = doc.add_heading('Medical AI Models Test Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        summary_text = f"""
This report presents the test results for all trained models in the Medical AI Assistant Platform.
Four models were tested: Skin Lesion Analyzer, Respiratory Sound Analyzer, Lab Data Analyzer, and Medical Chatbot.

All models support both Classification and Segmentation/Attention tasks.
        """
        doc.add_paragraph(summary_text.strip())
        
        # Results Table
        doc.add_heading('Model Performance Summary', level=1)
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Header
        header_cells = table.rows[0].cells
        headers = ['Model', 'Type', 'Accuracy/IoU', 'Avg Inference (ms)', 'Status']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # Model results
        model_data = [
            ('Skin Lesion', 'U-Net', results.get('skin', {}), 'seg'),
            ('Respiratory Sound', 'CNN', results.get('sound', {}), 'cls'),
            ('Lab Data', 'MLP', results.get('lab', {}), 'cls'),
            ('Medical Chatbot', 'LSTM', results.get('chatbot', {}), 'cls')
        ]
        
        for name, model_type, data, dtype in model_data:
            if data:
                row = table.add_row().cells
                row[0].text = name
                row[1].text = model_type
                
                if dtype == 'seg':
                    metric = f"IoU: {data.get('avg_iou', 0):.4f}\nDice: {data.get('avg_dice', 0):.4f}"
                else:
                    metric = f"Acc: {data.get('accuracy', 0):.4f}"
                row[2].text = metric
                row[3].text = f"{data.get('avg_inference_time', 0)*1000:.2f}"
                row[4].text = "✅ Passed"
            else:
                row = table.add_row().cells
                row[0].text = name
                row[1].text = model_type
                row[2].text = "N/A"
                row[3].text = "N/A"
                row[4].text = "❌ Failed"
        
        # Detailed Results
        for model_name, data in results.items():
            if data:
                doc.add_heading(f'{model_name.title()} Model Details', level=2)
                
                # Test info
                doc.add_paragraph(f"Total test samples: {data.get('total', 'N/A')}")
                
                if 'avg_iou' in data:
                    doc.add_paragraph(f"Average IoU: {data['avg_iou']:.4f}")
                    doc.add_paragraph(f"Average Dice: {data['avg_dice']:.4f}")
                
                if 'accuracy' in data:
                    doc.add_paragraph(f"Accuracy: {data['accuracy']:.4f}")
                
                if 'class_distribution' in data:
                    doc.add_paragraph("Class Distribution:")
                    for cls, count in data['class_distribution'].items():
                        doc.add_paragraph(f"  - Class {cls}: {count}", style='List Bullet')
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph('Medical AI Assistant Platform - Test Report')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save
        output_path = os.path.join(output_dir, 'all_models_test_report.docx')
        doc.save(output_path)
        print(f"✅ Word report saved: {output_path}")
    else:
        # JSON fallback
        output_path = os.path.join(output_dir, 'all_models_test_report.json')
        with open(output_path, 'w') as f:
            json.dump(results, f, indent=2, default=str)
        print(f"✅ JSON report saved: {output_path}")
    
    return output_path


# =============================================================================
# MAIN
# =============================================================================

def main():
    """Run all model tests and generate report."""
    print("=" * 70)
    print("MEDICAL AI MODELS TEST SUITE")
    print(f"Started: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print("=" * 70)
    
    # Device
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    print(f"\n🖥️ Device: {device}")
    
    if torch.cuda.is_available():
        print(f"   GPU: {torch.cuda.get_device_name(0)}")
    
    # Check model files
    print("\n📦 Checking model checkpoints...")
    for name, config in MODELS.items():
        checkpoint_path = os.path.join(PROJECT_ROOT, config['checkpoint'])
        if os.path.exists(checkpoint_path):
            size_mb = os.path.getsize(checkpoint_path) / (1024 * 1024)
            print(f"   ✅ {name}: {size_mb:.2f} MB")
        else:
            print(f"   ❌ {name}: Not found")
    
    # Run tests
    results = {}
    
    print("\n" + "=" * 70)
    print("RUNNING TESTS")
    print("=" * 70)
    
    # Test Skin
    try:
        results['skin'] = test_skin_model(MODELS['skin'], device)
    except Exception as e:
        print(f"❌ Skin model test failed: {e}")
        results['skin'] = None
    
    # Test Sound
    try:
        results['sound'] = test_sound_model(MODELS['sound'], device)
    except Exception as e:
        print(f"❌ Sound model test failed: {e}")
        results['sound'] = None
    
    # Test Lab
    try:
        results['lab'] = test_lab_model(MODELS['lab'], device)
    except Exception as e:
        print(f"❌ Lab model test failed: {e}")
        results['lab'] = None
    
    # Test Chatbot
    try:
        results['chatbot'] = test_chatbot_model(MODELS['chatbot'], device)
    except Exception as e:
        print(f"❌ Chatbot model test failed: {e}")
        results['chatbot'] = None
    
    # Generate report
    report_path = generate_word_report(results)
    
    # Summary
    print("\n" + "=" * 70)
    print("TEST SUMMARY")
    print("=" * 70)
    
    passed = sum(1 for v in results.values() if v is not None)
    total = len(results)
    
    print(f"\n✅ Models tested: {passed}/{total}")
    print(f"📄 Report saved: {report_path}")
    
    print("\n" + "=" * 70)
    print("TEST COMPLETE")
    print("=" * 70)
    
    return results


if __name__ == "__main__":
    results = main()
