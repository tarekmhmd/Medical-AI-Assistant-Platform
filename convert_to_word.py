"""
Convert Markdown Research Paper to Word Document
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_cell_border(cell, **kwargs):
    """Set cell border for tables"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def create_word_document():
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # Title
    title = doc.add_heading('A Comprehensive Analysis of the Medical Assistant Diagnostic Platform: Architecture, Implementation, and Applications', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Abstract
    doc.add_heading('Abstract', level=1)
    abstract_text = """This research paper presents a comprehensive analysis of the "first_version" GitHub repository, a medical assistant diagnostic platform designed to provide multi-modal health analysis capabilities. The project integrates artificial intelligence models for skin disease detection, respiratory sound analysis, and laboratory result interpretation, combined with an intelligent chatbot interface. This paper examines the technical architecture, implementation details, data management strategies, and potential applications of the system, while also discussing its strengths, limitations, and future development directions."""
    doc.add_paragraph(abstract_text)
    
    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    
    doc.add_heading('1.1 Project Overview', level=2)
    doc.add_paragraph("""The Medical Assistant Diagnostic Platform represents an innovative approach to healthcare technology, combining multiple AI-powered diagnostic tools into a unified web application. The project aims to democratize access to preliminary health assessments through image analysis, audio processing, and laboratory data interpretation. The system is designed to assist healthcare professionals and provide educational resources for patients seeking to understand their health conditions.""")
    
    doc.add_heading('1.2 Research Objectives', level=2)
    doc.add_paragraph('The primary objectives of this research are to:')
    objectives = [
        'Document the technical architecture and implementation details of the medical assistant platform',
        'Analyze the integration of multiple AI diagnostic modules',
        'Evaluate the project structure, file organization, and data management strategies',
        'Assess potential applications and limitations of the system',
        'Provide recommendations for future development and research'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    # 2. Project Architecture and Structure
    doc.add_heading('2. Project Architecture and Structure', level=1)
    
    doc.add_heading('2.1 Directory Structure', level=2)
    doc.add_paragraph('The project follows a well-organized modular architecture, separating frontend, backend, data, and model components:')
    
    # Directory structure table
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    headers = table.rows[0].cells
    headers[0].text = 'Directory'
    headers[1].text = 'Description'
    
    dirs = [
        ('backend/', 'Python backend application with database, models, and utilities'),
        ('frontend/', 'Web interface with HTML, CSS, and JavaScript files'),
        ('data/', 'Data storage for diseases, lab results, respiratory sounds, and skin images'),
        ('models_pretrained/', 'Pre-trained AI models (.h5 files)'),
        ('tessdata/', 'Tesseract OCR language data files'),
        ('tesseract/', 'Tesseract installation files'),
        ('doc/', 'Documentation files')
    ]
    for i, (dir_name, desc) in enumerate(dirs, 1):
        row = table.rows[i].cells
        row[0].text = dir_name
        row[1].text = desc
    
    doc.add_paragraph()
    
    doc.add_heading('2.2 Main Components', level=2)
    
    # Components table
    table2 = doc.add_table(rows=6, cols=3)
    table2.style = 'Table Grid'
    headers2 = table2.rows[0].cells
    headers2[0].text = 'Component'
    headers2[1].text = 'Location'
    headers2[2].text = 'Description'
    
    components = [
        ('Skin Analyzer', 'backend/models/skin_analyzer.py', 'Dermatological image analysis'),
        ('Sound Analyzer', 'backend/models/sound_analyzer.py', 'Respiratory sound classification'),
        ('Lab Analyzer', 'backend/models/lab_analyzer.py', 'Laboratory result interpretation'),
        ('Chatbot', 'backend/models/chatbot.py', 'AI-powered conversational interface'),
        ('Web Interface', 'frontend/', 'HTML/CSS/JS user interface')
    ]
    for i, (comp, loc, desc) in enumerate(components, 1):
        row = table2.rows[i].cells
        row[0].text = comp
        row[1].text = loc
        row[2].text = desc
    
    doc.add_paragraph()
    
    # 3. Technical Implementation
    doc.add_heading('3. Technical Implementation', level=1)
    
    doc.add_heading('3.1 Programming Languages and Frameworks', level=2)
    
    doc.add_paragraph('Backend Technologies:', style='Heading 3')
    backend_tech = [
        'Python 3.14.3 - Primary backend language',
        'TensorFlow/Keras - Deep learning framework',
        'Tesseract OCR - Optical character recognition for lab reports',
        'SQLite - Database management'
    ]
    for tech in backend_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_paragraph('Frontend Technologies:', style='Heading 3')
    frontend_tech = [
        'HTML5 - Web page structure',
        'CSS3 - Styling and responsive design',
        'JavaScript - Client-side interactivity'
    ]
    for tech in frontend_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_paragraph('Infrastructure:', style='Heading 3')
    infra_tech = [
        'Docker - Containerization',
        'Git LFS - Large file management',
        'Git - Version control'
    ]
    for tech in infra_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_heading('3.2 AI Model Architecture', level=2)
    doc.add_paragraph('The system incorporates three specialized deep learning models stored in .h5 format:')
    
    models = [
        'Skin Model - For dermatological image classification',
        'Sound Model - For respiratory sound pattern recognition',
        'Lab Model - For laboratory result analysis'
    ]
    for model in models:
        doc.add_paragraph(model, style='List Bullet')
    
    doc.add_heading('3.3 Git LFS Configuration', level=2)
    doc.add_paragraph('The project employs Git Large File Storage (LFS) to manage binary files efficiently:')
    
    # LFS table
    table3 = doc.add_table(rows=6, cols=2)
    table3.style = 'Table Grid'
    headers3 = table3.rows[0].cells
    headers3[0].text = 'File Type'
    headers3[1].text = 'Purpose'
    
    lfs_files = [
        ('*.dll', 'System libraries'),
        ('*.exe', 'Executable files'),
        ('*.traineddata', 'Tesseract OCR language models'),
        ('*.jar', 'Java archive files'),
        ('*.h5', 'Deep learning model files')
    ]
    for i, (ftype, purpose) in enumerate(lfs_files, 1):
        row = table3.rows[i].cells
        row[0].text = ftype
        row[1].text = purpose
    
    doc.add_paragraph()
    
    # 4. Data and Large Files Analysis
    doc.add_heading('4. Data and Large Files Analysis', level=1)
    
    doc.add_heading('4.1 Large File Statistics', level=2)
    
    # Large files table
    table4 = doc.add_table(rows=6, cols=4)
    table4.style = 'Table Grid'
    headers4 = table4.rows[0].cells
    headers4[0].text = 'File Type'
    headers4[1].text = 'Number of Files'
    headers4[2].text = 'Estimated Size'
    headers4[3].text = 'Purpose'
    
    large_files = [
        ('DLL files', '58', '~50-100 MB', 'System libraries'),
        ('EXE files', '18', '~20-50 MB', 'Executable binaries'),
        ('Trained data', '100+', '~1-2 GB', 'Tesseract OCR languages'),
        ('JAR files', '3', '~5 MB', 'Java dependencies'),
        ('H5 models', '3', '~100-500 MB', 'Pre-trained AI models')
    ]
    for i, (ftype, num, size, purpose) in enumerate(large_files, 1):
        row = table4.rows[i].cells
        row[0].text = ftype
        row[1].text = num
        row[2].text = size
        row[3].text = purpose
    
    doc.add_paragraph()
    
    doc.add_heading('4.2 Dataset Organization', level=2)
    doc.add_paragraph('The data directory structure indicates specialized storage for different medical data types:')
    
    data_dirs = [
        'diseases/ - Disease information database',
        'lab_results/ - Laboratory test results',
        'respiratory_sounds/ - Audio recordings for analysis',
        'skin_images/ - Dermatological images'
    ]
    for d in data_dirs:
        doc.add_paragraph(d, style='List Bullet')
    
    doc.add_heading('4.3 OCR Language Support', level=2)
    doc.add_paragraph('The tessdata directory contains extensive multilingual OCR support, enabling laboratory report processing in numerous languages including Arabic, English, Chinese Simplified/Traditional, and 100+ additional languages.')
    
    # 5. Functional Modules
    doc.add_heading('5. Functional Modules', level=1)
    
    doc.add_heading('5.1 Skin Disease Analysis', level=2)
    skin_features = [
        'Image upload and preprocessing',
        'Deep learning-based skin lesion classification',
        'Visual feedback and diagnostic suggestions'
    ]
    for f in skin_features:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_heading('5.2 Respiratory Sound Analysis', level=2)
    sound_features = [
        'Audio recording and processing',
        'Respiratory pattern recognition',
        'Abnormality detection in breathing sounds'
    ]
    for f in sound_features:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_heading('5.3 Laboratory Result Interpretation', level=2)
    lab_features = [
        'OCR-based report digitization',
        'Automatic result interpretation',
        'Reference range comparison',
        'Health recommendations'
    ]
    for f in lab_features:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_heading('5.4 Intelligent Chatbot', level=2)
    chat_features = [
        'Natural language health queries',
        'Contextual medical information',
        'Integration with diagnostic modules'
    ]
    for f in chat_features:
        doc.add_paragraph(f, style='List Bullet')
    
    # 6. Use Cases and Applications
    doc.add_heading('6. Use Cases and Applications', level=1)
    
    doc.add_heading('6.1 Clinical Applications', level=2)
    
    # Clinical applications table
    table5 = doc.add_table(rows=5, cols=3)
    table5.style = 'Table Grid'
    headers5 = table5.rows[0].cells
    headers5[0].text = 'Application Area'
    headers5[1].text = 'Module'
    headers5[2].text = 'Benefit'
    
    apps = [
        ('Dermatology', 'Skin Analyzer', 'Preliminary skin condition assessment'),
        ('Pulmonology', 'Sound Analyzer', 'Respiratory condition screening'),
        ('General Medicine', 'Lab Analyzer', 'Automated lab report interpretation'),
        ('Patient Education', 'Chatbot', 'Health information dissemination')
    ]
    for i, (area, module, benefit) in enumerate(apps, 1):
        row = table5.rows[i].cells
        row[0].text = area
        row[1].text = module
        row[2].text = benefit
    
    doc.add_paragraph()
    
    doc.add_heading('6.2 Research Applications', level=2)
    research_apps = [
        'Medical AI Development - Benchmark dataset for model training',
        'Telemedicine Research - Remote diagnostic capabilities',
        'Healthcare Accessibility - Democratizing medical knowledge',
        'Multi-modal Fusion - Integration of diverse diagnostic inputs'
    ]
    for app in research_apps:
        doc.add_paragraph(app, style='List Bullet')
    
    # 7. Evaluation and Limitations
    doc.add_heading('7. Evaluation and Limitations', level=1)
    
    doc.add_heading('7.1 Strengths', level=2)
    strengths = [
        'Multi-modal Integration - Combines three distinct diagnostic modalities',
        'Modular Architecture - Facilitates independent module development',
        'Containerization - Docker support ensures deployment consistency',
        'Multilingual Support - Extensive OCR language capabilities',
        'Comprehensive Documentation - Multiple markdown documentation files'
    ]
    for i, s in enumerate(strengths, 1):
        doc.add_paragraph(f'{i}. {s}')
    
    doc.add_heading('7.2 Limitations', level=2)
    limitations = [
        'Large Repository Size - ~3.6 GB total, requiring Git LFS management',
        'Model Transparency - Pre-trained models without training documentation',
        'Dependency Management - Multiple binary dependencies increase complexity',
        'Internet Dependency - Requires connectivity for full functionality',
        'Clinical Validation - Absence of clinical trial data'
    ]
    for i, l in enumerate(limitations, 1):
        doc.add_paragraph(f'{i}. {l}')
    
    doc.add_heading('7.3 Recommendations for Improvement', level=2)
    
    doc.add_paragraph('Short-term:', style='Heading 3')
    short_term = ['Add API documentation', 'Include model training scripts', 'Implement unit tests']
    for s in short_term:
        doc.add_paragraph(s, style='List Bullet')
    
    doc.add_paragraph('Medium-term:', style='Heading 3')
    medium_term = ['Clinical validation studies', 'Performance optimization', 'Mobile application development']
    for s in medium_term:
        doc.add_paragraph(s, style='List Bullet')
    
    doc.add_paragraph('Long-term:', style='Heading 3')
    long_term = ['Regulatory compliance (FDA, CE marking)', 'Multi-center clinical trials', 'Integration with EHR systems']
    for s in long_term:
        doc.add_paragraph(s, style='List Bullet')
    
    # 8. File Statistics and Visualizations
    doc.add_heading('8. File Statistics and Visualizations', level=1)
    
    doc.add_heading('8.1 Tracked File Distribution', level=2)
    
    # File distribution table
    table6 = doc.add_table(rows=9, cols=3)
    table6.style = 'Table Grid'
    headers6 = table6.rows[0].cells
    headers6[0].text = 'File Type'
    headers6[1].text = 'Count'
    headers6[2].text = 'Percentage'
    
    file_dist = [
        ('Markdown (.md)', '20', '23%'),
        ('Python (.py)', '14', '16%'),
        ('HTML (.html)', '12', '14%'),
        ('Batch (.bat)', '8', '9%'),
        ('CSS (.css)', '1', '1%'),
        ('JavaScript (.js)', '1', '1%'),
        ('Config files', '3', '3%'),
        ('Other', '27', '33%')
    ]
    for i, (ftype, count, pct) in enumerate(file_dist, 1):
        row = table6.rows[i].cells
        row[0].text = ftype
        row[1].text = count
        row[2].text = pct
    
    doc.add_paragraph()
    
    doc.add_heading('8.2 Documentation Coverage', level=2)
    
    # Documentation table
    table7 = doc.add_table(rows=7, cols=2)
    table7.style = 'Table Grid'
    headers7 = table7.rows[0].cells
    headers7[0].text = 'Document'
    headers7[1].text = 'Purpose'
    
    docs_list = [
        ('README.md', 'Project overview'),
        ('INSTALLATION_GUIDE.md', 'Installation instructions'),
        ('DEVELOPER_GUIDE.md', 'Development guidelines'),
        ('USER_MANUAL.md', 'End-user documentation'),
        ('AI_MODELS_GUIDE.md', 'AI model documentation'),
        ('DATA_SOURCES.md', 'Data provenance')
    ]
    for i, (doc_name, purpose) in enumerate(docs_list, 1):
        row = table7.rows[i].cells
        row[0].text = doc_name
        row[1].text = purpose
    
    doc.add_paragraph()
    
    # 9. Conclusion
    doc.add_heading('9. Conclusion', level=1)
    
    doc.add_heading('9.1 Summary of Contributions', level=2)
    doc.add_paragraph('The Medical Assistant Diagnostic Platform represents a significant contribution to healthcare technology, offering:')
    
    contributions = [
        'Integrated Diagnostic System - A unified platform combining skin, respiratory, and laboratory analysis',
        'AI-Powered Analysis - Deep learning models for automated diagnostic assistance',
        'Accessible Interface - Web-based design for broad accessibility',
        'Multilingual Support - Extensive OCR capabilities for global deployment',
        'Modular Design - Flexible architecture for future expansion'
    ]
    for i, c in enumerate(contributions, 1):
        doc.add_paragraph(f'{i}. {c}')
    
    doc.add_heading('9.2 Future Directions', level=2)
    doc.add_paragraph('The project presents numerous opportunities for future research and development:')
    
    future = [
        'Clinical Integration - Partnership with healthcare institutions for real-world validation',
        'Model Enhancement - Continuous improvement of diagnostic accuracy',
        'Regulatory Pathway - Pursuit of medical device certification',
        'Scale Deployment - Cloud-based deployment for broader access',
        'Research Collaboration - Open-source community engagement'
    ]
    for f in future:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_heading('9.3 Final Remarks', level=2)
    doc.add_paragraph("""This research paper has provided a comprehensive analysis of the Medical Assistant Diagnostic Platform, documenting its architecture, implementation, and potential applications. The project demonstrates the potential of AI-assisted diagnostic tools in healthcare, while also highlighting the challenges and considerations necessary for clinical deployment. Future work should focus on clinical validation, regulatory compliance, and expanding the platform's capabilities to serve diverse healthcare needs.""")
    
    # References
    doc.add_heading('References', level=1)
    refs = [
        'Project Repository: https://github.com/tarekmhmd/first_version',
        'Tesseract OCR Documentation: https://github.com/tesseract-ocr/tesseract',
        'TensorFlow/Keras Documentation: https://www.tensorflow.org/',
        'Git LFS Documentation: https://git-lfs.github.com/'
    ]
    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(f'[{i}] {ref}')
    
    doc.add_paragraph()
    doc.add_paragraph('Document generated: February 26, 2026')
    doc.add_paragraph('Repository version: first_version (Initial commit)')
    
    # Save document
    doc.save(r'D:\project 2\Research_Paper_Medical_Assistant_Platform.docx')
    print("Word document created successfully!")

if __name__ == "__main__":
    create_word_document()
